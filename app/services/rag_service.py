"""
RAG Service
Parse and embed documents for vector storage
"""

import os
import asyncio
import logging
from typing import List
import requests
from PyPDF2 import PdfReader
from docx import Document
# Updated import for newer LangChain version
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from google.cloud import firestore
from app.models.rag_document import RAGDocument
from app.database.vector_store import get_vector_store
from app.services.retriever_service import store_memory, get_embedding
from bs4 import BeautifulSoup
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Set up logging
logger = logging.getLogger(__name__)

class RAGService:
    """Service for parsing and embedding documents"""
    
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )
        google_api_key = os.getenv("GEMINI_API_KEY")
        
        try:
            self.embeddings = GoogleGenerativeAIEmbeddings(
                model="models/embedding-001",
                google_api_key=google_api_key
            )
        except Exception as e:
            logger.error(f"Failed to initialize embeddings: {e}")
            print(f"Failed to initialize embeddings: {e}")
            self.embeddings = None
            
        self._bg_tasks = {}  # Store background task status
    
    def get_task_status(self, task_id: str) -> dict:
        """Get status of a background task"""
        return self._bg_tasks.get(task_id, {"status": "not_found"})

    async def start_crawl_task(
        self,
        domain_url: str,
        campaign_id: str,
        db: firestore.Client,
        agent_id: str = None,
        max_pages: int = 50
    ) -> str:
        """Start a background crawl task and return task_id"""
        print(f"DEBUG: Starting crawl task for {domain_url}")
        import uuid
        try:
            task_id = str(uuid.uuid4())
            print(f"DEBUG: Generated task_id: {task_id}")
            
            # Initialize task status
            self._bg_tasks[task_id] = {
                "status": "pending",
                "progress": 0,
                "message": "Initializing crawler...",
                "details": {}
            }
            print(f"DEBUG: Task status initialized")
            
            # Start background task
            asyncio.create_task(self.process_domain(
                domain_url, campaign_id, db, agent_id, max_pages, task_id
            ))
            print(f"DEBUG: Background task created")
            
            return task_id
        except Exception as e:
            print(f"DEBUG: Error in start_crawl_task: {e}")
            import traceback
            traceback.print_exc()
            raise e

    async def process_document(
        self, 
        file_path: str, 
        file_type: str, 
        campaign_id: str,
        db: firestore.Client,
        agent_id: str = None
    ) -> RAGDocument:
        """
        Process a document and store its chunks in vector store
        """
        try:
            # Extract text based on file type
            content = await self._extract_text(file_path, file_type)
            
            # Split into chunks
            chunks = self.text_splitter.split_text(content)
            
            # Create RAG document record
            rag_doc = RAGDocument(
                campaign_id=str(campaign_id) if campaign_id else None,
                agent_id=str(agent_id) if agent_id else None,
                filename=os.path.basename(file_path),
                content=content,
                file_type=file_type,
                chunks_extracted=len(chunks)
            )
            
            # Add to Firestore
            update_time, doc_ref = db.collection('rag_documents').add(rag_doc.to_dict())
            rag_doc.id = doc_ref.id
            
            # Embed and store chunks
            await self._embed_and_store_chunks(chunks, rag_doc.id, campaign_id, agent_id)
            
            return rag_doc
            
        except Exception as e:
            print(f"Error in process_document: {e}")
            # No rollback needed for Firestore, but we might want to delete the doc if created
            raise e
    
    async def process_url(
        self, 
        url: str, 
        campaign_id: str,
        db: firestore.Client,
        agent_id: str = None
    ) -> RAGDocument:
        """
        Process a URL and store its content in vector store
        """
        try:
            # Fetch content from URL
            response = requests.get(url, timeout=30)
            response.raise_for_status()
            
            # Parse HTML content
            soup = BeautifulSoup(response.content, 'html.parser')
            for script in soup(["script", "style"]):
                script.decompose()
            content = soup.get_text()
            
            # Clean up text
            lines = (line.strip() for line in content.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            content = ' '.join(chunk for chunk in chunks if chunk)
            
            # Split into chunks
            text_chunks = self.text_splitter.split_text(content)
            
            # Get title
            title = soup.title.string if soup.title else url
            
            # Create RAG document record
            rag_doc = RAGDocument(
                campaign_id=str(campaign_id) if campaign_id else None,
                agent_id=str(agent_id) if agent_id else None,
                filename=url,
                title=title,
                content=content,
                file_type="url",
                chunks_extracted=len(text_chunks)
            )
            
            # Add to Firestore
            update_time, doc_ref = db.collection('rag_documents').add(rag_doc.to_dict())
            rag_doc.id = doc_ref.id
            
            # Embed and store chunks
            await self._embed_and_store_chunks(text_chunks, rag_doc.id, campaign_id, agent_id)
            
            return rag_doc
            
        except Exception as e:
            print(f"Error in process_url: {e}")
            raise e
    
    async def process_domain(
        self,
        domain_url: str,
        campaign_id: str,
        db: firestore.Client,
        agent_id: str = None,
        max_pages: int = 50,
        task_id: str = None
    ) -> dict:
        """
        Process an entire domain by crawling all pages
        """
        from app.services.web_scraper import WebScraper
        
        try:
            if task_id:
                self._bg_tasks[task_id].update({
                    "status": "crawling",
                    "progress": 0,
                    "message": f"Crawling {domain_url}..."
                })

            scraper = WebScraper(max_pages=max_pages, delay=3.0, concurrency=3)
            
            # Callback for scraper progress
            async def scraper_progress(scraped, total):
                if task_id:
                    # Crawling is first 50% of progress
                    progress = min(45, int((scraped / max_pages) * 45))
                    self._bg_tasks[task_id].update({
                        "progress": progress,
                        "message": f"Crawling: {scraped} pages found found so far..."
                    })
            
            result = await scraper.scrape_website(domain_url, progress_callback=scraper_progress)
            
            if task_id:
                 self._bg_tasks[task_id].update({
                    "status": "processing",
                    "progress": 50,
                    "message": f"Processing {len(result['content'])} pages..."
                })
            
            documents = []
            failed_urls = result['failed_urls']
            
            semaphore = asyncio.Semaphore(5)
            
            total_pages = len(result['content'])
            processed_count = 0
            
            async def process_page(page_data):
                nonlocal processed_count
                async with semaphore:
                    try:
                        text_chunks = self.text_splitter.split_text(page_data['content'])
                        
                        rag_doc = RAGDocument(
                            campaign_id=str(campaign_id) if campaign_id else None,
                            agent_id=str(agent_id) if agent_id else None,
                            filename=page_data['url'],
                            title=page_data['title'],
                            content=page_data['content'],
                            file_type="url",
                            chunks_extracted=len(text_chunks)
                        )
                        
                        update_time, doc_ref = db.collection('rag_documents').add(rag_doc.to_dict())
                        rag_doc.id = doc_ref.id
                        
                        await self._embed_and_store_chunks(text_chunks, rag_doc.id, campaign_id, agent_id)
                        
                        processed_count += 1
                        if task_id and total_pages > 0:
                            # Processing is 50-95%
                            progress = 50 + int((processed_count / total_pages) * 45)
                            self._bg_tasks[task_id].update({
                                "progress": progress,
                                "message": f"Processed {processed_count}/{total_pages} pages"
                            })
                        
                        return rag_doc
                    except Exception as e:
                        print(f"Error processing page {page_data['url']}: {e}")
                        failed_urls.append({
                            'url': page_data['url'],
                            'reason': str(e)
                        })
                        return None

            tasks = [process_page(page_data) for page_data in result['content']]
            processed_docs = await asyncio.gather(*tasks)
            documents = [doc for doc in processed_docs if doc is not None]
                    
            return {
                'documents': documents,
                'total_pages': result['total_pages'],
                'failed_urls': failed_urls
            }
            
        except Exception as e:
            if task_id:
                self._bg_tasks[task_id].update({
                    "status": "failed",
                    "error": str(e)
                })
            print(f"Error in process_domain: {e}")
            raise e
        
        if task_id:
            self._bg_tasks[task_id].update({
                "status": "completed",
                "progress": 100,
                "message": "Crawling and processing completed successfully!",
                "result": {
                    'total_pages': result['total_pages'],
                    'documents_count': len(documents),
                    'failed_count': len(failed_urls)
                }
            })
    
    async def _extract_text(self, file_path: str, file_type: str) -> str:
        """Extract text from file based on type"""
        logger.info(f"Extracting text from file: {file_path}, file_type: {file_type}")
        if file_type.lower() == "pdf":
            return self._extract_pdf_text(file_path)
        elif file_type.lower() == "docx":
            return self._extract_docx_text(file_path)
        else:
            supported_types = ["pdf", "docx"]
            import os
            normalized_path = file_path.replace('\\', '/')
            file_extension = os.path.splitext(normalized_path)[1].lower()
            raise ValueError(f"Unsupported file type: {file_path}. Supported file types are: {', '.join(supported_types)}")
    
    def _extract_pdf_text(self, file_path: str) -> str:
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text()
        return text
    
    def _extract_docx_text(self, file_path: str) -> str:
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    
    async def _embed_and_store_chunks(
        self, 
        chunks: List[str], 
        document_id: str, 
        campaign_id: str,
        agent_id: str = None
    ):
        """Embed text chunks and store in vector store"""
        try:
            embeddings = self.embeddings.embed_documents(chunks)
            
            documents = []
            document_ids = []
            for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                doc_id = f"{document_id}_{i}"
                document_ids.append(doc_id)
                
                metadata = {
                    "document_id": document_id,
                    "campaign_id": campaign_id,
                    "agent_id": agent_id,
                    "chunk_index": i
                }
                metadata = {k: v for k, v in metadata.items() if v is not None}

                documents.append({
                    "id": doc_id,
                    "text": chunk,
                    "embedding": embedding,
                    "metadata": metadata
                })
                
                client_id = str(agent_id) if agent_id else str(campaign_id)
                
                memory_metadata = {
                    "document_id": document_id,
                    "campaign_id": campaign_id,
                    "agent_id": agent_id,
                    "chunk_index": i,
                    "source": "rag_document"
                }
                memory_metadata = {k: v for k, v in memory_metadata.items() if v is not None}
                
                store_memory(client_id, chunk, memory_metadata)
            
            vector_store = get_vector_store()
            doc_embeddings = [doc["embedding"] for doc in documents]
            vector_store.add_documents(doc_embeddings, document_ids)
            
        except Exception as e:
            if "quota exceeded" in str(e).lower() or "429" in str(e):
                logger.warning("Google API quota exceeded, using fallback embedding method")
                await self._embed_and_store_chunks_fallback(chunks, document_id, campaign_id)
            else:
                print(f"Error embedding and storing chunks: {e}")
                raise e
    
    async def _embed_and_store_chunks_fallback(
        self, 
        chunks: List[str], 
        document_id: str, 
        campaign_id: str,
        agent_id: str = None
    ):
        """Fallback method to embed text chunks using mock embeddings"""
        try:
            documents = []
            document_ids = []
            for i, chunk in enumerate(chunks):
                embedding = get_embedding(chunk)
                doc_id = f"{document_id}_{i}"
                document_ids.append(doc_id)
                
                metadata = {
                    "document_id": document_id,
                    "campaign_id": campaign_id,
                    "agent_id": agent_id,
                    "chunk_index": i,
                    "source": "rag_document",
                    "fallback_embedding": True
                }
                metadata = {k: v for k, v in metadata.items() if v is not None}

                documents.append({
                    "id": doc_id,
                    "text": chunk,
                    "embedding": embedding,
                    "metadata": metadata
                })
                
                client_id = str(agent_id) if agent_id else str(campaign_id)
                
                memory_metadata = {
                    "document_id": document_id,
                    "campaign_id": campaign_id,
                    "agent_id": agent_id,
                    "chunk_index": i,
                    "source": "rag_document",
                    "fallback_embedding": True
                }
                memory_metadata = {k: v for k, v in memory_metadata.items() if v is not None}
                
                store_memory(client_id, chunk, memory_metadata)
            
            zero_embeddings = [[0.0] * 768 for _ in chunks]
            
            vector_store = get_vector_store()
            vector_store.add_documents(zero_embeddings, document_ids)
            
        except Exception as e:
            print(f"Error in fallback embedding method: {e}")
            raise e


# Global instance (Lazy loaded)
_rag_service_instance = None

def get_rag_service():
    global _rag_service_instance
    if _rag_service_instance is None:
        _rag_service_instance = RAGService()
    return _rag_service_instance

# rag_service = RAGService() # Removed synchronous init